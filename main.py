import tkinter as tk
from tkinter import messagebox
import customtkinter as ctk

import docx
from tkinter import filedialog
import time
import os
import sys

if hasattr(sys, '_MEIPASS'):
    # PyInstallerで1ファイル化された場合に解凍される一時フォルダのパス
    base_dir = sys._MEIPASS
else:
    # 通常のPythonスクリプトとして実行した場合のパス
    base_dir = os.path.dirname(os.path.abspath(__file__))

ctk.set_default_color_theme("blue")

va1 = ["購入", "使用"]

font_body = ("FOT-筑紫B丸ゴシック Std R", 16)
main_bg_color = "aliceblue"
main_fg_color = "white"


class App(ctk.CTk):
    def __init__(self):
        super().__init__()

        # 関連変数の定義
        self.flag_mode = 0
        self.flag_use = 1
        self.ward_use = ""

        # UI構築とイベントループの開始
        self._build_ui()
        self.mainloop()

    def button_end(self):
        me_end = tk.messagebox.askyesno("終了の確認", "プログラムを終了しますか？")
        if me_end:
            self.destroy()

    def button_enter(self):
        if self.flag_use == 1:
            self.ward_use = "購入"
        elif self.flag_use == 2:
            self.ward_use = "使用"
        else:
            tk.messagebox.showerror("エラー", "コンボボックス内をすべて揃えてください。")
            return

        # 使用するPDMを指定
        doc = docx.Document(os.path.join(base_dir, "PDM_ARF-001.docx"))

        # 会計要望書の追記
        # 日付出力
        doc.paragraphs[0].text = self.en0_1.get()
        # 団体名出力
        doc.paragraphs[5].text = "東京電機大学東京千住キャンパス" + self.en0_2.get()
        # 代表者名出力
        doc.paragraphs[7].text = "代表　　" + self.en0_3_1.get() + "　" + self.en0_3_2.get()
        # 会計担当者名出力
        doc.paragraphs[8].text = "会計　　" + self.en0_4_1.get() + "　" + self.en0_4_2.get()
        checkbox_paragraph_map = [
            (self.cb0_1, 12),
            (self.cb0_2, 13),
            (self.cb0_3, 14),
            (self.cb0_4, 15),
            (self.cb0_5, 16),
            (self.cb0_6, 17),
            (self.cb0_7, 18)
        ]
        for cb, i in checkbox_paragraph_map:
            if cb.get() == 1:
                t = doc.paragraphs[i].text
                t = t.replace("□", "☑")
                doc.paragraphs[i].text = t

        # 要件出力
        doc.paragraphs[20].text = self.te0_1.get("1.0", "end-1c")

        if self.flag_mode == 0:
            # 品名出力
            doc.paragraphs[25].text = "　　　　　　　　　　　品名　　　" + self.en1_1.get()
            # 金額出力
            doc.paragraphs[26].text = "　　　　　　　　　　　金額　　　\\" + self.en1_2.get() + "-"
            # 購入予定日出力
            doc.paragraphs[27].text = "　　　　　　　　　" + self.ward_use + "予定日　　" + self.en1_3.get()
            # 購入理由出力
            doc.paragraphs[28].text = "　　　　　　　　　" + self.ward_use + "理由　　　" + self.te1_1.get("1.0", "end-1c")
        elif self.flag_mode == 1:
            # 見積金額（往復）出力
            doc.paragraphs[25].text = "　　　　見積金額（往復）　　　　\\" + self.en2_1.get() + "-"
            # 使用月日出力
            doc.paragraphs[26].text = "　　　　使用月日　　　　　　　　" + self.en2_2.get()
            # 使用経路（往復）出力
            doc.paragraphs[27].text = "　　　　使用理由　　　　　　　　" + self.te2_1.get("1.0", "end-1c")
            doc.paragraphs[28].text = "　　　　使用経路（往復）　　　　" + self.en2_3.get()
            doc.paragraphs[29].insert_paragraph_before("　　　　　　　　　　　　　　　　" + self.en2_4.get())
            doc.paragraphs[29].insert_paragraph_before("　　　　　　　　　　　　　　　　" + self.en2_5.get())
            doc.paragraphs[29].insert_paragraph_before("　　　　　　　　　　　　　　　　" + self.en2_6.get())
            doc.paragraphs[29].insert_paragraph_before("　　　　　　　　　　　　　　　　" + self.en2_7.get())
        elif self.flag_mode == 2:
            # 使用交通機関
            doc.paragraphs[25].text = "　　　　　　　　　使用交通機関　　　" + self.en3_1.get()
            # 区間出力
            doc.paragraphs[26].text = "　　　　　　　　　区間　　　　　　　" + self.en3_2.get()
            # 金額出力
            doc.paragraphs[27].text = "　　　　　　　　　金額　　　　　　　" + self.en3_3.get() + "円"
            # 内訳出力
            doc.paragraphs[28].text = "　　　　　　　　　内訳　　　　　　　" + self.en3_4.get()
            # 購入予定日出力
            doc.paragraphs[29].insert_paragraph_before("　　　　　　　　　購入予定日　　　　" + self.en3_5.get())
            # 購入理由
            doc.paragraphs[30].insert_paragraph_before("　　　　　　　　　購入理由　　　　　" + self.te3_1.get("1.0", "end-1c"))
        elif self.flag_mode == 3:
            # 車種出力
            doc.paragraphs[25].text = "　　　　　　　　　　　車種　　　　　" + self.en4_1.get()
            # 見積金額出力
            doc.paragraphs[26].text = "　　　　　　　　　見積金額　　　　　\\" + self.en4_2.get() + "-"
            # ガソリン代出力
            doc.paragraphs[27].text = "　　　　　　　　　ガソリン代　　　　\\" + self.en4_3.get() + "-"
            # 高速道路代出力
            doc.paragraphs[28].text = "　　　　　　　　　高速道路代　　　　\\" + self.en4_4.get() + "-"
            # 使用予定日出力
            doc.paragraphs[29].insert_paragraph_before("　　　　　　　　　使用予定日　　　　" + self.en4_5.get())
            # 使用経路出力
            doc.paragraphs[30].insert_paragraph_before("　　　　　　　　　使用経路（往復）　" + self.en4_6.get())
            # 使用理由
            doc.paragraphs[31].insert_paragraph_before("　　　　　　　　　使用理由　　　　　" + self.te4_1.get("1.0", "end-1c"))
        else:
            tk.messagebox.showerror("エラー", "エラーが発生しました。設定を変更してください。")
            return

        # 文書の保存
        file_name = filedialog.asksaveasfilename(title="作成する文書の保存",
                                                 initialfile=time.strftime("%Y_%m%d_") + "会計要望書",
                                                 defaultextension=".docx",
                                                 filetypes=[("Word 文書", ".docx")])
        if file_name:
            doc.save(file_name)
            # 保存した文書を開く（.docxを開く設定を行っているアプリケーションで）
            os.startfile(file_name)
        
        self.destroy()

    # ウィンドウの様式を変更する関数（汎用仕様）
    def button_change1(self):
        self.flag_mode = 0
        self.te0_1.delete("1.0", "end")
        self.te0_1.insert("1.0", "○○代を予算内から使用することを要望いたします。")
        self.bu_change1.configure(fg_color="black", hover_color="black", state="disabled")
        self.bu_change2.configure(fg_color="white", hover_color="whitesmoke", text_color="black", state="normal")
        self.bu_change3.configure(fg_color="white", hover_color="whitesmoke", text_color="black", state="normal")

        self.fr_mode1.grid(column=0, row=16, columnspan=3, padx=20, pady=10)
        self.fr_mode2.grid_forget()
        self.fr_mode3.grid_forget()

        self.cb0_1.configure(state=tk.NORMAL)
        self.cb0_2.configure(state=tk.DISABLED)
        self.cb0_3.configure(state=tk.NORMAL)
        self.cb0_4.configure(state=tk.NORMAL)
        self.cb0_5.configure(state=tk.NORMAL)
        self.cb0_6.configure(state=tk.NORMAL)

        self.cb0_2.deselect()

    # ウィンドウの様式を変更する関数（交通費（電車・バス）仕様）
    def button_change2(self):
        self.flag_mode = 1
        self.te0_1.delete("1.0", "end")
        self.te0_1.insert("1.0", "○○への移動として交通費を予算内から使用することを要望いたします。")
        self.bu_change1.configure(fg_color="white", hover_color="whitesmoke", text_color="black", state="normal")
        self.bu_change2.configure(fg_color="black", hover_color="black", state="disabled")
        self.bu_change3.configure(fg_color="white", hover_color="whitesmoke", state="normal")
        self.bu_change4.configure(fg_color="white", hover_color="whitesmoke", state="normal")

        self.switch_mode_frame(self.fr_mode2)
        self.set_checkboxes_for_transport()

    # ウィンドウの様式を変更する関数（交通費（電車・バス／複数経路）仕様）
    def button_change3(self):
        self.flag_mode = 2
        self.te0_1.delete("1.0", "end")
        self.te0_1.insert("1.0", "○○への移動として交通費を予算内から使用することを要望いたします。")
        self.bu_change1.configure(fg_color="white", hover_color="whitesmoke", text_color="black", state="normal")
        self.bu_change2.configure(fg_color="white", hover_color="whitesmoke", state="normal")
        self.bu_change3.configure(fg_color="black", hover_color="black", state="disabled")
        self.bu_change4.configure(fg_color="white", hover_color="whitesmoke", state="normal")

        self.switch_mode_frame(self.fr_mode3)
        self.set_checkboxes_for_transport()

    # ウィンドウの様式を変更する関数（交通費（レンタカー）仕様）
    def button_change4(self):
        self.flag_mode = 3
        self.te0_1.delete("1.0", "end")
        self.te0_1.insert("1.0", "○○のため以下の要領でレンタカーを使用することを要望いたします。")
        self.bu_change1.configure(fg_color="white", hover_color="whitesmoke", text_color="black", state="normal")
        self.bu_change2.configure(fg_color="white", hover_color="whitesmoke", state="normal")
        self.bu_change3.configure(fg_color="white", hover_color="whitesmoke", state="normal")
        self.bu_change4.configure(fg_color="black", hover_color="black", state="disabled")

        self.switch_mode_frame(self.fr_mode4)
        self.set_checkboxes_for_transport()

    # モードフレームを切り替えるヘルパー関数
    def switch_mode_frame(self, active_frame):
        for fr in (self.fr_mode1, self.fr_mode2, self.fr_mode3, self.fr_mode4):
            fr.grid_forget()
        active_frame.grid(column=0, row=16, columnspan=3, padx=20, pady=10)

    # 交通費モード時のチェックボックス一括設定ヘルパー関数
    def set_checkboxes_for_transport(self):
        for cb in (self.cb0_1, self.cb0_2, self.cb0_3, self.cb0_4, self.cb0_5, self.cb0_6):
            cb.configure(state=tk.DISABLED)
            cb.deselect()
        self.cb0_2.select()

    # コンボボックスを選択すると起動する関数
    def combo_select(self, e):
        indices = [va1.index(cb.get()) if cb.get() in va1 else -1
                   for cb in (self.co1_1, self.co1_2, self.co1_3, self.co1_4)]
        if all(i == 0 for i in indices):
            self.flag_use = 1
        elif all(i == 1 for i in indices):
            self.flag_use = 2
        else:
            self.flag_use = 0

    # ウィンドウ上でマウスホイールを回すとスクロールバーが移動する関数
    def mouse_y_scroll(self, event):
        if event.delta < 0:
            self.canvas.yview_scroll(1, "units")
        else:
            self.canvas.yview_scroll(-1, "units")

    def _build_ui(self):
        # 汎用会計要望書入力要求ウィンドウの仕様設定
        self.title("会計要望書生成")
        self.geometry("800x800")
        self.resizable(False, True)
        self.bind("<MouseWheel>", self.mouse_y_scroll)

        # スクロールバー関連
        frame = tk.Frame(self, bg=main_bg_color)
        frame.pack(fill=tk.BOTH, expand=True)

        self.canvas = tk.Canvas(frame, bg=main_bg_color)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scrollbar_y = tk.Scrollbar(frame, orient=tk.VERTICAL, command=self.canvas.yview)
        scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
        self.canvas.configure(yscrollcommand=scrollbar_y.set)

        scrollable_frame = ctk.CTkFrame(self.canvas, fg_color=main_bg_color, width=800, height=1350)
        self.canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")

        scrollable_frame.update_idletasks()
        self.canvas.config(scrollregion=self.canvas.bbox("all"))

        # タイトル
        la_title = ctk.CTkLabel(scrollable_frame, text="会計要望書生成", font=("FOT-筑紫B丸ゴシック Std R", 60))
        la_title.grid(column=0, row=0, columnspan=3, padx=20, pady=10)

        # 入力関連ウィジェット（共通部分）を配置するフレーム
        fr_body = ctk.CTkFrame(scrollable_frame, fg_color=main_fg_color, width=600, corner_radius=10)
        fr_body.grid(column=0, row=1, columnspan=3, padx=20, pady=10)

        # モード切替ボタンを配置するフレーム
        la_button1 = ctk.CTkLabel(fr_body, text="モードを選択してください（黒背景が選択中モード）。", font=font_body)
        la_button1.grid(column=0, row=14, sticky="w", padx=20, pady=10)

        fr_button1 = ctk.CTkFrame(fr_body, fg_color=main_fg_color, width=800, height=50)
        fr_button1.grid(column=0, row=15, columnspan=3, padx=20, pady=20)

        # 入力関連ウィジェット（汎用仕様部分）を配置するフレーム
        self.fr_mode1 = ctk.CTkFrame(fr_body, fg_color=main_fg_color)
        self.fr_mode1.grid(column=0, row=16, columnspan=3, padx=20, pady=10)

        # 入力関連ウィジェット（交通費（複数経路）仕様部分）を配置するフレーム
        self.fr_mode2 = ctk.CTkFrame(fr_body, fg_color=main_fg_color)

        # 入力関連ウィジェット（交通費（単一経路）仕様部分）を配置するフレーム
        self.fr_mode3 = ctk.CTkFrame(fr_body, fg_color=main_fg_color)

        # 入力関連ウィジェット（レンタカー仕様部分）を配置するフレーム
        self.fr_mode4 = ctk.CTkFrame(fr_body, fg_color=main_fg_color)

        # 入力ボタンと終了ボタンを配置するフレーム
        fr_button2 = ctk.CTkFrame(scrollable_frame, fg_color=main_bg_color, width=800, height=50)
        fr_button2.grid(column=0, row=50, columnspan=3, padx=20, pady=10)

        # fr_bodyに配置する入力関連ウィジェットの作成
        la0_1 = ctk.CTkLabel(fr_body, text="要望書を提出する日付を入力してください。", font=font_body)
        self.en0_1 = ctk.CTkEntry(fr_body, width=200, placeholder_text="例：令和6年1月1日")
        la0_2 = ctk.CTkLabel(fr_body, text="団体名を入力してください。", font=font_body)
        self.en0_2 = ctk.CTkEntry(fr_body, width=200, placeholder_text="例：新聞委員会")
        la0_3 = ctk.CTkLabel(fr_body, text="代表者名を入力してください。", font=font_body)
        self.en0_3_1 = ctk.CTkEntry(fr_body, width=100, placeholder_text="代表者の名字")
        self.en0_3_2 = ctk.CTkEntry(fr_body, width=100, placeholder_text="代表者の名前")
        la0_4 = ctk.CTkLabel(fr_body, text="会計担当者名を入力してください。", font=font_body)
        self.en0_4_1 = ctk.CTkEntry(fr_body, width=100, placeholder_text="担当者の名字")
        self.en0_4_2 = ctk.CTkEntry(fr_body, width=100, placeholder_text="担当者の名前")
        la0_5 = ctk.CTkLabel(fr_body, text="要望書の用件を以下から選んでください。", font=font_body)
        self.cb0_1 = ctk.CTkCheckBox(fr_body, text="遠征等で資材などの運搬を依頼する時", font=font_body)
        self.cb0_2 = ctk.CTkCheckBox(fr_body, text="交通機関を用いた移動費として使用する時（電車・バス等）", font=font_body, state=tk.DISABLED)
        self.cb0_3 = ctk.CTkCheckBox(fr_body, text="領収書単位で税抜き５万円を超えている場合", font=font_body)
        self.cb0_4 = ctk.CTkCheckBox(fr_body, text="銀行振り込み等で領収証が発行できないとき恐れがある場合", font=font_body)
        self.cb0_5 = ctk.CTkCheckBox(fr_body, text="三部会五委員会が郵送代・飲食費を交際費として使用したい場合", font=font_body)
        self.cb0_6 = ctk.CTkCheckBox(fr_body, text="自治会費で購入してよいか迷ったとき", font=font_body)
        self.cb0_7 = ctk.CTkCheckBox(fr_body, text="その他", font=font_body)
        la0_6 = ctk.CTkLabel(fr_body, text="要望内容を書き換えてください（例：○○代⇒カメラ代）。", font=font_body)
        self.te0_1 = ctk.CTkTextbox(fr_body, width=650, height=75, font=font_body)
        self.te0_1.insert("1.0", "○○代を予算内から使用することを要望いたします。")

        # fr_mode1に配置する入力関連ウィジェットの作成
        self.co1_1 = ctk.CTkComboBox(self.fr_mode1, state="readonly", values=va1, font=font_body, width=75, command=self.combo_select)
        la1_1 = ctk.CTkLabel(self.fr_mode1, text="する物の品名を入力してください。", font=font_body)
        self.en1_1 = ctk.CTkEntry(self.fr_mode1, width=200, placeholder_text="例：カメラ")
        self.co1_2 = ctk.CTkComboBox(self.fr_mode1, state="readonly", values=va1, font=font_body, width=75, command=self.combo_select)
        la1_2 = ctk.CTkLabel(self.fr_mode1, text="する物の合計金額を入力してください。", font=font_body)
        self.en1_2 = ctk.CTkEntry(self.fr_mode1, width=200, placeholder_text="例：350,000")
        self.co1_3 = ctk.CTkComboBox(self.fr_mode1, state="readonly", values=va1, font=font_body, width=75, command=self.combo_select)
        la1_3 = ctk.CTkLabel(self.fr_mode1, text="予定日を入力してください。", font=font_body)
        self.en1_3 = ctk.CTkEntry(self.fr_mode1, width=200, placeholder_text="例：1月31日")
        self.co1_4 = ctk.CTkComboBox(self.fr_mode1, state="readonly", values=va1, font=font_body, width=75, command=self.combo_select)
        la1_4 = ctk.CTkLabel(self.fr_mode1, text="する理由を入力してください。", font=font_body)
        self.te1_1 = ctk.CTkTextbox(self.fr_mode1, width=600, height=60, font=font_body)

        # fr_mode1のコンボボックス関連設定
        self.co1_1.set(va1[0])
        self.co1_2.set(va1[0])
        self.co1_3.set(va1[0])
        self.co1_4.set(va1[0])

        # fr_mode2に配置する入力関連ウィジェットの作成
        la2_1 = ctk.CTkLabel(self.fr_mode2, text="往復分の合計見積金額を入力してください。", font=font_body)
        self.en2_1 = ctk.CTkEntry(self.fr_mode2, width=200, placeholder_text="例：10,000")
        la2_2 = ctk.CTkLabel(self.fr_mode2, text="使用予定日を入力してください。", font=font_body)
        self.en2_2 = ctk.CTkEntry(self.fr_mode2, width=200, placeholder_text="例：1月31日")
        la2_3 = ctk.CTkLabel(self.fr_mode2, text="使用理由を入力してください。", font=font_body)
        self.te2_1 = ctk.CTkTextbox(self.fr_mode2, width=600, height=60, font=font_body)
        la2_4 = ctk.CTkLabel(self.fr_mode2, text="往復分の乗車経路を入力してください。", font=font_body)
        self.en2_3 = ctk.CTkEntry(self.fr_mode2, width=280, placeholder_text="例：北千住→日暮里→池袋→高坂")
        self.en2_4 = ctk.CTkEntry(self.fr_mode2, width=280)
        self.en2_5 = ctk.CTkEntry(self.fr_mode2, width=280)
        self.en2_6 = ctk.CTkEntry(self.fr_mode2, width=280)
        self.en2_7 = ctk.CTkEntry(self.fr_mode2, width=600)

        # fr_mode3に配置する入力関連ウィジェットの作成
        la3_1 = ctk.CTkLabel(self.fr_mode3, text="使用する交通機関名（路線名等）を入力してください。", font=font_body)
        self.en3_1 = ctk.CTkEntry(self.fr_mode3, width=250, placeholder_text="例：常磐線、山手線、東武東上線")
        la3_2 = ctk.CTkLabel(self.fr_mode3, text="乗車区間を入力してください。", font=font_body)
        self.en3_2 = ctk.CTkEntry(self.fr_mode3, width=250, placeholder_text="例：北千住→日暮里→池袋→高坂")
        la3_3 = ctk.CTkLabel(self.fr_mode3, text="要望する合計金額を入力してください。", font=font_body)
        self.en3_3 = ctk.CTkEntry(self.fr_mode3, width=250, placeholder_text="例：10,000")
        la3_4 = ctk.CTkLabel(self.fr_mode3, text="合計金額の内訳を入力してください。", font=font_body)
        self.en3_4 = ctk.CTkEntry(self.fr_mode3, width=250, placeholder_text="例：往復1,000円　×10名")
        la3_5 = ctk.CTkLabel(self.fr_mode3, text="使用予定日を入力してください。", font=font_body)
        self.en3_5 = ctk.CTkEntry(self.fr_mode3, width=250, placeholder_text="例：1月31日")
        la3_6 = ctk.CTkLabel(self.fr_mode3, text="使用理由を入力してください。", font=font_body)
        self.te3_1 = ctk.CTkTextbox(self.fr_mode3, width=600, height=60, font=font_body)

        # fr_mode4に配置する入力関連ウィジェットの作成
        la4_1 = ctk.CTkLabel(self.fr_mode4, text="使用する車種を入力してください。", font=font_body)
        self.en4_1 = ctk.CTkEntry(self.fr_mode4, width=250, placeholder_text="例：乗用自動車")
        la4_2 = ctk.CTkLabel(self.fr_mode4, text="合計見積金額を入力してください。", font=font_body)
        self.en4_2 = ctk.CTkEntry(self.fr_mode4, width=250, placeholder_text="例：50,000")
        la4_3 = ctk.CTkLabel(self.fr_mode4, text="ガソリン代の見積金額を入力してください。", font=font_body)
        self.en4_3 = ctk.CTkEntry(self.fr_mode4, width=250, placeholder_text="例：20,000")
        la4_4 = ctk.CTkLabel(self.fr_mode4, text="高速道路代の見積金額を入力してください。", font=font_body)
        self.en4_4 = ctk.CTkEntry(self.fr_mode4, width=250, placeholder_text="例：5,000")
        la4_5 = ctk.CTkLabel(self.fr_mode4, text="使用予定日を入力してください。", font=font_body)
        self.en4_5 = ctk.CTkEntry(self.fr_mode4, width=250, placeholder_text="例：1月31日")
        la4_6 = ctk.CTkLabel(self.fr_mode4, text="乗車経路（往復）を入力してください。", font=font_body)
        self.en4_6 = ctk.CTkEntry(self.fr_mode4, width=600, placeholder_text="例：北千住→東京電機大学埼玉鳩山キャンパス")
        la4_7 = ctk.CTkLabel(self.fr_mode4, text="使用理由を入力してください。", font=font_body)
        self.te4_1 = ctk.CTkTextbox(self.fr_mode4, width=600, height=60, font=font_body)

        # fr_bodyに配置する入力関連ウィジェットの配置
        la0_1.grid(column=0, row=0, sticky="w", padx=20, pady=10)
        self.en0_1.grid(column=1, row=0, columnspan=2, padx=20, pady=10)
        la0_2.grid(column=0, row=1, sticky="w", padx=20, pady=10)
        self.en0_2.grid(column=1, row=1, columnspan=2, padx=20, pady=10)
        la0_3.grid(column=0, row=2, sticky="w", padx=20, pady=10)
        self.en0_3_1.grid(column=1, row=2, padx=0, pady=10)
        self.en0_3_2.grid(column=2, row=2, padx=20, pady=10)
        la0_4.grid(column=0, row=3, sticky="w", padx=20, pady=10)
        self.en0_4_1.grid(column=1, row=3, padx=0, pady=10)
        self.en0_4_2.grid(column=2, row=3, padx=20, pady=10)
        la0_5.grid(column=0, row=4, sticky="w", columnspan=3, padx=20, pady=10)
        self.cb0_1.grid(column=0, row=5, sticky="w", columnspan=3, padx=50, pady=5)
        self.cb0_2.grid(column=0, row=6, sticky="w", columnspan=3, padx=50, pady=5)
        self.cb0_3.grid(column=0, row=7, sticky="w", columnspan=3, padx=50, pady=5)
        self.cb0_4.grid(column=0, row=8, sticky="w", columnspan=3, padx=50, pady=5)
        self.cb0_5.grid(column=0, row=9, sticky="w", columnspan=3, padx=50, pady=5)
        self.cb0_6.grid(column=0, row=10, sticky="w", columnspan=3, padx=50, pady=5)
        self.cb0_7.grid(column=0, row=11, sticky="w", columnspan=3, padx=50, pady=5)
        la0_6.grid(column=0, row=12, sticky="w", padx=20, pady=10)
        self.te0_1.grid(column=0, row=13, sticky="w", columnspan=3, padx=20, pady=10)

        # fr_mode1に配置する入力関連ウィジェットの配置
        self.co1_1.grid(column=0, row=0, sticky="w", pady=10)
        la1_1.grid(column=1, row=0, sticky="w", pady=10)
        self.en1_1.grid(column=2, row=0, sticky="e", pady=10)
        self.co1_2.grid(column=0, row=1, sticky="w", pady=10)
        la1_2.grid(column=1, row=1, sticky="w", pady=10)
        self.en1_2.grid(column=2, row=1, sticky="e", pady=10)
        self.co1_3.grid(column=0, row=2, sticky="w", pady=10)
        la1_3.grid(column=1, row=2, sticky="w", pady=10)
        self.en1_3.grid(column=2, row=2, sticky="e", pady=10)
        self.co1_4.grid(column=0, row=3, sticky="w", pady=10)
        la1_4.grid(column=1, row=3, sticky="w", pady=10)
        self.te1_1.grid(column=0, row=4, sticky="w", columnspan=3, pady=10)

        # fr_mode2に配置する入力関連ウィジェットの配置
        la2_1.grid(column=0, row=0, sticky="w", pady=10)
        self.en2_1.grid(column=1, row=0, sticky="e", pady=10)
        la2_2.grid(column=0, row=1, sticky="w", pady=10)
        self.en2_2.grid(column=1, row=1, sticky="e", pady=10)
        la2_3.grid(column=0, row=2, sticky="w", pady=10)
        self.te2_1.grid(column=0, row=3, sticky="w", columnspan=2, pady=5)
        la2_4.grid(column=0, row=4, sticky="w", pady=5)
        self.en2_3.grid(column=0, row=5, sticky="w", pady=5)
        self.en2_4.grid(column=1, row=5, sticky="e", pady=5)
        self.en2_5.grid(column=0, row=6, sticky="w", pady=5)
        self.en2_6.grid(column=1, row=6, sticky="e", pady=5)
        self.en2_7.grid(column=0, row=7, sticky="w", columnspan=2, pady=5)

        # fr_mode3に配置する入力関連ウィジェットの配置
        la3_1.grid(column=0, row=0, sticky="w", pady=5)
        self.en3_1.grid(column=1, row=0, sticky="e", pady=5)
        la3_2.grid(column=0, row=1, sticky="w", pady=5)
        self.en3_2.grid(column=1, row=1, sticky="e", pady=5)
        la3_3.grid(column=0, row=2, sticky="w", pady=5)
        self.en3_3.grid(column=1, row=2, sticky="e", pady=5)
        la3_4.grid(column=0, row=3, sticky="w", pady=5)
        self.en3_4.grid(column=1, row=3, sticky="e", pady=5)
        la3_5.grid(column=0, row=4, sticky="w", pady=5)
        self.en3_5.grid(column=1, row=4, sticky="e", pady=5)
        la3_6.grid(column=0, row=5, sticky="w", pady=5)
        self.te3_1.grid(column=0, row=6, sticky="w", columnspan=2, pady=5)

        # fr_mode4に配置する入力関連ウィジェットの配置
        la4_1.grid(column=0, row=0, sticky="w", pady=5)
        self.en4_1.grid(column=1, row=0, sticky="e", pady=5)
        la4_2.grid(column=0, row=1, sticky="w", pady=5)
        self.en4_2.grid(column=1, row=1, sticky="e", pady=5)
        la4_3.grid(column=0, row=2, sticky="w", pady=5)
        self.en4_3.grid(column=1, row=2, sticky="e", pady=5)
        la4_4.grid(column=0, row=3, sticky="w", pady=5)
        self.en4_4.grid(column=1, row=3, sticky="e", pady=5)
        la4_5.grid(column=0, row=4, sticky="w", pady=5)
        self.en4_5.grid(column=1, row=4, sticky="e", pady=5)
        la4_6.grid(column=0, row=5, sticky="w", pady=5)
        self.en4_6.grid(column=0, row=6, sticky="e", columnspan=2, pady=5)
        la4_7.grid(column=0, row=7, sticky="w", pady=5)
        self.te4_1.grid(column=0, row=8, sticky="w", columnspan=2, pady=5)

        # 各種ボタンの作成
        self.bu_change1 = ctk.CTkButton(fr_button1, text="汎用",
                                   font=("游ゴシック", 20, "bold"), text_color="white",
                                   width=70, height=40, corner_radius=0, border_width=1,
                                   fg_color="black", hover_color="black", text_color_disabled="white",
                                   command=self.button_change1, state="disabled")
        self.bu_change2 = ctk.CTkButton(fr_button1, text="交通費（複数経路）",
                                   font=("游ゴシック", 20, "bold"), text_color="black",
                                   width=200, height=40, corner_radius=0, border_width=1,
                                   fg_color="white", hover_color="whitesmoke", text_color_disabled="white",
                                   command=self.button_change2)
        self.bu_change3 = ctk.CTkButton(fr_button1, text="交通費（同一経路）",
                                   font=("游ゴシック", 20, "bold"), text_color="black",
                                   width=200, height=40, corner_radius=0, border_width=1,
                                   fg_color="white", hover_color="whitesmoke", text_color_disabled="white",
                                   command=self.button_change3)
        self.bu_change4 = ctk.CTkButton(fr_button1, text="レンタカー",
                                   font=("游ゴシック", 20, "bold"), text_color="black",
                                   width=150, height=40, corner_radius=0,  border_width=1,
                                   fg_color="white", hover_color="whitesmoke", text_color_disabled="white",
                                   command=self.button_change4)

        bu_enter = ctk.CTkButton(fr_button2, text="入力",
                                 font=("游ゴシック", 20, "bold"), text_color="white",
                                 width=70, height=40,
                                 fg_color="mediumseagreen", hover_color="seagreen",
                                 command=self.button_enter)
        bu_end = ctk.CTkButton(fr_button2, text="終了",
                               font=("游ゴシック", 20, "bold"), text_color="white",
                               width=70, height=40,
                               fg_color="red", hover_color="crimson",
                               command=self.button_end)

        # 各種ボタンの配置
        self.bu_change1.grid(column=0, row=0, padx=10)
        self.bu_change2.grid(column=1, row=0, padx=10)
        self.bu_change3.grid(column=2, row=0, padx=10)
        self.bu_change4.grid(column=3, row=0, padx=10)
        bu_enter.grid(column=0, row=0, padx=10)
        bu_end.grid(column=1, row=0, padx=10)

if __name__ == "__main__":
    App()
